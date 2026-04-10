"""RAG: extract PDF/DOCX, chunk, embed once with sentence-transformers, retrieve for Questions."""

from __future__ import annotations

import json
import re
import uuid
from io import BytesIO
from typing import Any

import numpy as np
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import RagChunk, RagDocument

# Lazy singleton
_model: Any = None

CHUNK_SIZE = 900
CHUNK_OVERLAP = 120
TOP_K = 6
# In "auto" mode: use RAG if best cosine similarity >= this (all-MiniLM-L6-v2)
SIMILARITY_THRESHOLD = 0.28


def _get_embedder() -> Any:
    global _model
    if _model is None:
        from sentence_transformers import SentenceTransformer

        name = (settings.rag_embedding_model or "all-MiniLM-L6-v2").strip()
        _model = SentenceTransformer(name)
    return _model


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def extract_text(filename: str, mime: str, data: bytes) -> str:
    m = (mime or "").lower()
    fn = filename.lower()
    if "pdf" in m or fn.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if "word" in m or "officedocument" in m or fn.endswith(".docx"):
        return extract_text_from_docx(data)
    raise ValueError("Unsupported file type. Use PDF or DOCX.")


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    chunks: list[str] = []
    i = 0
    while i < len(text):
        piece = text[i : i + size]
        if piece.strip():
            chunks.append(piece.strip())
        i += size - overlap
        if i <= 0:
            i = size
    return chunks


def embed_texts_sync(texts: list[str]) -> np.ndarray:
    if not texts:
        return np.array([])
    model = _get_embedder()
    emb = model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    return np.asarray(emb, dtype=np.float32)


def embedding_to_json(vec: np.ndarray) -> str:
    return json.dumps(vec.astype(float).tolist())


def json_to_embedding(s: str) -> np.ndarray:
    return np.asarray(json.loads(s), dtype=np.float32)


def cosine_scores(query_vec: np.ndarray, matrix: np.ndarray) -> np.ndarray:
    """matrix: (n, dim)"""
    qn = query_vec / (np.linalg.norm(query_vec) + 1e-9)
    mn = matrix / (np.linalg.norm(matrix, axis=1, keepdims=True) + 1e-9)
    return mn @ qn


def keyword_match(question: str, keywords_csv: str | None) -> bool:
    if not keywords_csv or not keywords_csv.strip():
        return False
    q = question.lower()
    for k in keywords_csv.split(","):
        k = k.strip().lower()
        if len(k) >= 2 and k in q:
            return True
    return False


def topic_overlap(question: str, topic_tags: str | None) -> bool:
    """True if any topic tag word appears in the question (for auto mode document filter)."""
    if not topic_tags or not topic_tags.strip():
        return True
    q = question.lower()
    for t in topic_tags.replace(";", ",").split(","):
        t = t.strip().lower()
        if len(t) >= 2 and t in q:
            return True
    return False


async def ingest_document(
    db: AsyncSession,
    *,
    filename: str,
    mime: str,
    data: bytes,
    topic_tags: str | None,
) -> RagDocument:
    text = extract_text(filename, mime, data)
    if not text:
        raise ValueError("No extractable text in document.")
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("Document produced no chunks after processing.")

    embeddings = embed_texts_sync(chunks)
    if len(embeddings) != len(chunks):
        raise ValueError("Embedding batch size mismatch.")

    doc_id = str(uuid.uuid4())
    doc = RagDocument(
        id=doc_id,
        filename=filename[:500],
        mime=mime[:120],
        topic_tags=(topic_tags or "").strip() or None,
    )
    db.add(doc)
    for idx, (chunk, row) in enumerate(zip(chunks, embeddings)):
        db.add(
            RagChunk(
                document_id=doc_id,
                chunk_index=idx,
                content=chunk[:50_000],
                embedding_json=embedding_to_json(row),
            )
        )
    await db.commit()
    await db.refresh(doc)
    return doc


async def delete_document(db: AsyncSession, doc_id: str) -> bool:
    await db.execute(delete(RagChunk).where(RagChunk.document_id == doc_id))
    r = await db.execute(delete(RagDocument).where(RagDocument.id == doc_id))
    await db.commit()
    return (r.rowcount or 0) > 0


async def list_documents(db: AsyncSession) -> list[dict[str, Any]]:
    q = (
        select(RagDocument.id, RagDocument.filename, RagDocument.topic_tags, RagDocument.created_at)
        .order_by(RagDocument.created_at.desc())
    )
    rows = (await db.execute(q)).all()
    out: list[dict[str, Any]] = []
    for rid, fn, tags, created in rows:
        cnt = await db.scalar(
            select(func.count()).select_from(RagChunk).where(RagChunk.document_id == rid)
        )
        out.append(
            {
                "id": rid,
                "filename": fn,
                "topic_tags": tags,
                "chunk_count": int(cnt or 0),
                "created_at": created.isoformat() if created else "",
            }
        )
    return out


def retrieve_context_sync(
    question: str,
    rows: list[tuple[str, str, str, str]],
    *,
    always_include_best: bool,
) -> tuple[str | None, float]:
    """
    rows: (chunk_content, embedding_json, topic_tags, document_id)
    If always_include_best (RAG mode \"on\"), take top-K chunks regardless of score.
    Otherwise prefer chunks above a similarity floor.
    """
    if not rows or not question.strip():
        return None, 0.0
    embs: list[np.ndarray] = []
    contents: list[str] = []
    for content, ej, _tags, _did in rows:
        contents.append(content)
        embs.append(json_to_embedding(ej))
    mat = np.stack(embs, axis=0)
    qv = embed_texts_sync([question.strip()])[0]
    scores = cosine_scores(qv, mat)
    best = float(scores.max()) if scores.size else 0.0
    order = np.argsort(-scores)
    top_idx = order[:TOP_K]
    picked: list[str] = []
    if always_include_best:
        picked = [contents[int(i)] for i in top_idx]
    else:
        floor = SIMILARITY_THRESHOLD * 0.85
        picked = [contents[int(i)] for i in top_idx if scores[int(i)] >= floor]
        if not picked and len(top_idx):
            picked = [contents[int(top_idx[0])]]
    if not picked:
        return None, best
    block = "\n\n---\n\n".join(picked)
    return block[:12000], best


async def load_chunk_rows_for_retrieval(
    db: AsyncSession,
    question: str,
    *,
    rag_mode: str,
) -> list[tuple[str, str, str, str]]:
    """
    In \"auto\" mode, exclude chunks from documents that have topic_tags set unless
    the question overlaps those topics (documents with no tags are always searchable).
    """
    q = select(RagChunk.content, RagChunk.embedding_json, RagDocument.topic_tags, RagDocument.id).join(
        RagDocument, RagChunk.document_id == RagDocument.id
    )
    res = await db.execute(q)
    all_rows = res.all()
    auto = (rag_mode or "auto").strip().lower() == "auto"
    out: list[tuple[str, str, str, str]] = []
    for content, ej, tags, did in all_rows:
        if auto and tags and tags.strip() and not topic_overlap(question, tags):
            continue
        out.append((content, ej, tags or "", did))
    return out


def should_use_rag(
    *,
    mode: str,
    question: str,
    rag_keywords: str | None,
    has_documents: bool,
    best_similarity: float,
) -> bool:
    if not has_documents:
        return False
    m = (mode or "auto").strip().lower()
    if m == "off":
        return False
    if m == "on":
        return True
    # auto
    if keyword_match(question, rag_keywords):
        return True
    if best_similarity >= SIMILARITY_THRESHOLD:
        return True
    return False


def compute_rag_context_from_rows(
    retrieval_query: str,
    rag_mode: str,
    rag_keywords: str | None,
    rows: list[tuple[str, str, str, str]],
) -> str | None:
    """
    Sync: embed + score chunks. Call via run_in_threadpool after load_chunk_rows_for_retrieval.
    retrieval_query is the user question (answer) or recent transcript excerpt (suggest).
    """
    q = (retrieval_query or "").strip()
    if (rag_mode or "off").strip().lower() == "off" or not rows or not q:
        return None
    km = keyword_match(q, rag_keywords)
    always_best = rag_mode == "on" or (rag_mode == "auto" and km)
    ctx_rag, best_sim = retrieve_context_sync(
        q,
        rows,
        always_include_best=always_best,
    )
    use = should_use_rag(
        mode=rag_mode,
        question=q,
        rag_keywords=rag_keywords,
        has_documents=True,
        best_similarity=best_sim,
    )
    if use and ctx_rag:
        return ctx_rag
    return None
